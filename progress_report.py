"""
Takes the desired columns from formsite csv file,
and puts pre-determined info into docx file where the students are separated by teams
with bullet lists of the PM, members, added up scores, and comments.

Some names will need to be manually added after making the doc. Most of the time
the PM will put the name in the team member comment section.

Also some scores will also need to be manually add due to "?" being found
occasionally where numbers are supposed to be.
"""
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def docx_print(d):
    """
        Function takes needed information for docx file from dictionary
        and puts them in to a smaller dictionary separated by exectutive mentor name.
    """
    exec_d = {}
    for r in d:
        row_num = d[r]
        r_data = row_num["Row Data"]

        proj_title = "{} ({})".format(r_data["Project Title"], r_data["Project Type"])
        pm_n = "PM: {}".format(r_data["PM Name"])
        team = []

        for t in r_data["Team Members"]:

            t_name = t["Team Member Name"]
            _score_ = 0
            try:
                data_list = list(t["data"].values())
                data_list_int = [int(s) for s in data_list]
                _score_ += sum(data_list_int)
            except ValueError:
                _score_ = np.nan
            team.append({"name": str(t_name), "score": _score_})
        text = r_data["Comments"]
        if r_data["Executive"] in exec_d:
            exec_d[r_data["Executive"]].append({"proj_title": proj_title,
                "pm_n": pm_n,
                "team" : team,
                "comment": text})
        else:
            exec_d[r_data["Executive"]] = [{"proj_title": proj_title,
            "pm_n": pm_n,
            "team" : team,
            "comment": text}]
    return exec_d

document = Document()

reports = pd.read_csv("ProgressreportformCapstoneProg(4).csv") #Change to valid file name
report_dict = {}
num = int(input("Enter progress report num: "))
reports = reports.loc[reports['Progress Report #'] == num]
reports.to_csv("progress_report_output/progress_report_#{}.csv".format(str(num)),index=False)
print(reports["PM Name"])

""" Dictionary format
{
   Row #num': {
       Row # : num
       Row Data: {
           Project Type: Project Type,
           Project Title : Project Title,
           Executive: Name,
           PM Name : Name,
           Team Members : [
                {
                    Team Member Name: Team member name,
                    data: {
                       Overall: Score,
                       Professionalism: Score,
                       Technical: Score,
                       Communication: Score
                       Promptness: Score,
                       Ability to Get Along: Score,
                       Ability to Learn: Score
                    }
               },
               ...
          ]
          Comments : comment
       }
   }
}
"""
reports_cols = list(reports.columns)

project_type = [reports_cols[3]]
executive = [reports_cols[4]]
pm_name = [reports_cols[6]]
project_name = [reports_cols[10]]
tm_1_evals = reports_cols[11:20]
tm_2_evals = reports_cols[20:29]
tm_3_evals = reports_cols[29:38]
tm_4_evals = reports_cols[38:47]
tm_5_evals = reports_cols[47:56]
comments_col = ["Any obstacles or problems facing your project?"]
team_member_evals = tm_1_evals+tm_2_evals+tm_3_evals+tm_4_evals+tm_5_evals
working_cols = project_type + executive + pm_name + project_name + team_member_evals+comments_col

num_rows = reports.index.tolist()

tm_reports = reports.loc[:, working_cols]
print(tm_reports)

#drop_suffix = [" \(Item #31\)"," \(Item #34\)"," \(Item #38\)"," \(Item #41\)"," \(Item #44\)"]
#print(drop_suffix)
#for suf in drop_suffix:
#    tm_reports.columns = tm_reports.columns.str.replace(suf, "")

""" Change column names, previously used str.replace"""
for col in tm_reports.columns:
    index = col.find(" (Item")
    #print(index)
    if index == -1:
        continue
    v = col[:index]
    tm_reports.rename(columns={col: v}, inplace=True)

tm_reports.to_csv("progress_report_output/tm_reports.csv",index=False)

""" Makes hashmaps of each row, then populates the report data hashmap.
    May have to modify if Formsite format changes.
 """
for i in num_rows:
    row = tm_reports.loc[i]
    pm = row["PM Name"]
    row_dict = None
    project_type = row["What type of project do you have?"]
    project_title = row["Project Title"]
    executive = row["Executive Team Mentor"]
    tm_list = []
    for j in range(1,6):
        tm_dict = {}
        overall = None
        pro = None
        tech = None
        comm = None
        prompt = None
        g_along = None
        learn = None
        comment = None
        if j == 1:
            name = row["Team member name"]
            comment = row["Comments / Suggestions"]
        else:
            name = row["Team member name.{}".format(str(j-1))]
            comment = row["Comments / Suggestions.{}".format(str(j-1))]

        if j != 2:
            overall = row["Team Member#{} Overall".format(str(j))]
            pro = row["Team Member#{} Professionalism".format(str(j))]
            tech = row["Team Member#{} TechnicalSkills".format(str(j))]
            comm = row["Team Member#{} CommunicationSkills".format(str(j))]
            prompt = row["Team Member#{} Promptness".format(str(j))]
            g_along = row["Team Member#{} Ability toGet Along with Others".format(str(j))]
            learn = row["Team Member#{} Ability to Learn".format(str(j))]
        else:
            overall = row["Team Member # {} Overall".format(str(j))]
            pro = row["Team Member # {} Professionalism".format(str(j))]
            tech = row["Team Member # {} TechnicalSkills".format(str(j))]
            comm = row["Team Member # {} CommunicationSkills".format(str(j))]
            prompt = row["Team Member # {} Promptness".format(str(j))]
            g_along = row["Team Member # {} Ability toGet Along with Others".format(str(j))]
            learn = row["Team Member # {} Ability to Learn".format(str(j))]

        comment = row["Any obstacles or problems facing your project?"]
        tm_dict = {"Team Member Name": name,
                   "data" : {"Overall": overall,
                           "Professionalism": pro,
                            "Technical": tech,
                            "Communication": comm,
                            "Promptness": prompt,
                            "Ability to Get Along": g_along,
                            "Ability to Learn": learn,
                           }
                   }

        tm_list.append(tm_dict)
    row_dict = {
        "Row #" : i,
        "Row Data": {
            "Project Type" : project_type,
            "Project Title" : project_title,
            "PM Name" : pm,
            "Executive": executive,
            "Team Members" : tm_list,
            "Comments": comment
        }
    }
    report_dict["Row #{}".format(i)] = row_dict
#print(report_dict)

#exec_list = list(tm_reports["Exectutive Team Mentor"].unique())



""" Puts information into docx file in desired format"""
docx_dict = docx_print(report_dict)
exec_list = list(docx_dict.keys())
print(docx_dict.keys())
print(exec_list)
for name in exec_list:
    exec_title = document.add_paragraph()
    exec_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exec_runner = exec_title.add_run(name)
    exec_runner.bold = True
    exec_runner.underline = True
    font = exec_runner.font
    font.size = Pt(15)

    for proj in docx_dict[name]:
        p_title = proj["proj_title"]
        p_pm_name = proj["pm_n"]
        p_tm_title = "Team Members:"

        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p_bold = p.add_run(p_title)
        p_bold.bold = True

        p_pm_name = document.add_paragraph(p_pm_name, style='List Bullet 2')
        p_pm_name.paragraph_format.space_before = Pt(3)
        p_pm_name.paragraph_format.space_after = Pt(3)

        p_tm_title = document.add_paragraph(p_tm_title, style='List Bullet 2')
        p_tm_title.paragraph_format.space_before = Pt(3)
        p_tm_title.paragraph_format.space_after = Pt(3)

        for n in proj["team"]:
            tm_name = n["name"]
            score = n["score"]
            tm_name = tm_name.rsplit('\t', 1)[0]
            tm_name += " ({}/70)".format(str(score))
            tm_name = document.add_paragraph(tm_name, style='List Bullet 3')
            tm_name.paragraph_format.space_before = Pt(3)
            tm_name.paragraph_format.space_after = Pt(3)

        comment_title = document.add_paragraph("Comments/Issues", style='List Bullet 2')
        comment_title.paragraph_format.space_before = Pt(3)
        comment_title.paragraph_format.space_after = Pt(3)

        comment_text = proj["comment"]
        if proj["comment"] is np.nan:
            comment_text = "N/A"
        else:
            comment_text = proj["comment"]
        comments = document.add_paragraph(comment_text, style='List Bullet 3')
        comments.paragraph_format.space_before = Pt(3)
        comments.paragraph_format.space_after = Pt(3)
        #space = document.add_paragraph("\n")

document.save("progress_report_output/progress_report_#{}.docx".format(str(num)))
